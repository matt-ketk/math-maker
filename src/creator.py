from docx import Document
from lxml import etree
import os
from datetime import datetime
from simple_problem import SimpleProblem
import random

DIR = '../worksheets/'
COLS = 5

MML2OMML_PATH = 'C:/Program Files/Microsoft Office/Office16/MML2OMML.XSL'
MATHML_STRING = '<mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mfrac><mml:mrow><mml:mtable><mml:mtr><mml:mtd><mml:mo>{}</mml:mo></mml:mtd><mml:mtd><mml:mtable><mml:mtr><mml:mtd><mml:mi>{}</mml:mi></mml:mtd><mml:mtd><mml:mi>{}</mml:mi></mml:mtd></mml:mtr><mml:mtr><mml:mtd><mml:mi>{}</mml:mi></mml:mtd><mml:mtd><mml:mi>{}</mml:mi></mml:mtd></mml:mtr></mml:mtable></mml:mtd></mml:mtr></mml:mtable></mml:mrow><mml:mrow><mml:mtable><mml:mtr><mml:mtd></mml:mtd><mml:mtd></mml:mtd><mml:mtd></mml:mtd></mml:mtr></mml:mtable></mml:mrow></mml:mfrac></mml:math>'

MATH_DIC = {
    'addition' : '+',
    'subtraction' : '-'
}
# TIMESTAMP: datetime.now().strftime('%H:%M:%S')

def create_simple_sheet(name='additiontest2', type='addition', numproblems=50, min=1, max=99, dir=DIR):
    # det amnt of rows
    rows = 0
    if numproblems % COLS == 0:
        rows = int(numproblems / COLS)
    else:
        rows = 1 + int(numproblems / COLS)

    problem_set = [[0] * COLS for i in range(rows)]
    for r in range(rows):
        if (r != rows - 1):
            for c in range(COLS):
                problem_set[r][c] = SimpleProblem(a=random.randint(1,20),b=2,operation=type)
        else:
            for c in range(numproblems % COLS):
                problem_set[r][c] = SimpleProblem(a=random.randint(1,20),b=2,operation=type)
    # create doc
    d = Document()
    # create table, create var pointer
    problems = d.add_table(rows=rows, cols=COLS)
    # parse problem objs 
    r = 0
    c = 0
    for row in problem_set:
        # POTENTIAL SCALABILITY ISSUE: there is currently a maximum of two digits, not three nor four.
        for p in row:
            if p == 0:
                break
            a_tens_digit = str(int(p.a / 10))
            a_ones_digit = str(int(p.a % 10))

            b_tens_digit = str(int(p.b / 10))
            b_ones_digit = str(int(p.b % 10))

            if a_tens_digit == 0:
                a_tens_digit = ' '
            if b_tens_digit == 0:
                a_tens_digit = ' '
            mathml_string = MATHML_STRING.format(MATH_DIC[p.operation], a_tens_digit, a_ones_digit, b_tens_digit, b_ones_digit)
            tree = etree.fromstring(mathml_string)
            xslt = etree.parse(MML2OMML_PATH)

            transform = etree.XSLT(xslt)
            new_dom = transform(tree)
            cell = problems.cell(r, c)
            paragraph = cell.add_paragraph()
            paragraph._element.append(new_dom.getroot())
            if c == COLS - 1:
                r += 1
                c = 0
            else:
                c += 1
            

    d.save(DIR + name + '.docx')