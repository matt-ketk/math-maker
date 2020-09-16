from docx import Document
from lxml import etree


MML2OMML_PATH = 'C:/Program Files/Microsoft Office/Office16/MML2OMML.XSL'
# Convert MathML (MML) into Office MathML (OMML) using a XSLT stylesheet
mathml_string = '<mml:math xmlns:mml="http://www.w3.org/1998/Math/MathML" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><mml:mfrac><mml:mrow><mml:mtable><mml:mtr><mml:mtd><mml:mo>{}</mml:mo></mml:mtd><mml:mtd><mml:mtable><mml:mtr><mml:mtd><mml:mi>{}</mml:mi></mml:mtd><mml:mtd><mml:mi>{}</mml:mi></mml:mtd></mml:mtr><mml:mtr><mml:mtd><mml:mi>{}</mml:mi></mml:mtd><mml:mtd><mml:mi>d>{}</mml:mi></mml:mtd></mml:mtr></mml:mtable></mml:mtd></mml:mtr></mml:mtable></mml:mrow><mml:mrow><mml:mtable><mml:mtr><mml:mtd></mml:mtd><mml:mtd></mml:mtd><mml:mtd></mml:mtd></mml:mtr></mml:mtable></mml:mrow></mml:mfrac></mml:math>'
tree = etree.fromstring(mathml_string)
xslt = etree.parse(MML2OMML_PATH)

transform = etree.XSLT(xslt)
new_dom = transform(tree)

doc = Document('../test.docx')
"""
paragraph = doc.add_paragraph()
paragraph._element.append(new_dom.getroot())
"""
tbl = doc.add_table(5,5)
cell = tbl.cell(0,0)
cell.text = 'testing'
doc.save('../test.docx')

"""
CONVERTING OMML to MML:
1. in equations tab, select equations options (the arrow on the bottom left corner of the conversion section)
2. select "Copy MathML to clipboard as plain text"
3. open notepad and paste
"""